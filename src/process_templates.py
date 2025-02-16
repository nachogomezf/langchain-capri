import os
import json
from docx import Document
import re

def extract_sections_from_docx(doc_path):
    """Extract structured sections from a business plan DOCX file"""
    doc = Document(doc_path)
    current_section = None
    sections = {}
    content = []
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
            
        # Check if this is a section header (all caps or numbered)
        if text.isupper() or re.match(r'^\d+\.[\d\.]*\s+[A-Z]', text):
            current_section = text
            sections[current_section] = []
        elif current_section:
            sections[current_section].append(text)
        
        content.append(text)
    
    return sections, "\n".join(content)

def process_business_plan(doc_path):
    """Process a business plan document and extract structured information"""
    sections, full_text = extract_sections_from_docx(doc_path)
    
    # Try to extract key information
    investment_match = re.search(r'\$(\d+(?:,\d{3})*(?:\.\d{2})?[Kk]?)', full_text)
    investment = investment_match.group(0) if investment_match else "$500K"
    
    # Try to find business type from filename or content
    filename = os.path.basename(doc_path)
    business_type = filename.split('-')[-1].split('.')[0].strip()
    
    # Format the training example with sections
    formatted_text = ""
    for section, paragraphs in sections.items():
        formatted_text += f"{section}\n"
        formatted_text += "\n".join(paragraphs) + "\n\n"
    
    return {
        "input": f"Generate a business plan for a {investment} {business_type} E-2 Visa startup",
        "output": formatted_text
    }

def clean_question_text(text):
    """Clean up template artifacts from question text"""
    # Remove underscores and template lines
    text = re.sub(r'_+', '', text)
    # Remove multiple spaces
    text = re.sub(r'\s+', ' ', text)
    # Clean up any remaining template artifacts like blank lines or dots
    text = re.sub(r'\.{3,}', '', text)
    # Remove any remaining template placeholders
    text = re.sub(r'\([Yy]es/[Nn]o\)', '', text)
    return text.strip()

def determine_field_type(question_text):
    """Determine the appropriate field type based on question content"""
    question_lower = question_text.lower()
    
    # Financial questions
    if re.search(r'amount|budget|revenue|salary|investment|funds|costs?|expenses?|\$', question_lower):
        return "number", {"min": 0, "step": 1000}
    
    # Multi-line text fields
    if any(phrase in question_lower for phrase in [
        'describe', 'provide details', 'list', 'explain', 
        'breakdown', 'what are your', 'plans:', 'strategy'
    ]) or len(question_text) > 100:
        return "text_area", {}
    
    # Yes/No questions
    if re.search(r'\?.*\([Yy]es/[Nn]o\)|^[Dd]o you|[Ww]ill you|[Hh]ave you', question_lower):
        return "radio", {"options": ["Yes", "No"]}
    
    # Number questions
    if any(phrase in question_lower for phrase in [
        'how many', 'number of', 'quantity', 'count'
    ]):
        return "number", {"min": 0}
    
    return "text", {}

def extract_questionnaire_sections(doc_path):
    """Extract sections and questions from the E2 questionnaire"""
    doc = Document(doc_path)
    current_section = None
    sections = {}
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
            
        # If text is all caps or has a colon at the end, treat as section header
        if text.isupper() or text.endswith(':'):
            current_section = text.rstrip(':')
            sections[current_section] = {
                "title": current_section,
                "questions": []
            }
        elif current_section and '?' in text or ': ' in text:
            # Clean up the question text
            cleaned_text = clean_question_text(text)
            if cleaned_text:  # Only add if we have actual content
                # Generate a clean ID from the text
                question_id = re.sub(r'[^a-zA-Z0-9]', '_', cleaned_text.lower())
                question_id = re.sub(r'_+', '_', question_id)  # Replace multiple underscores with single
                question_id = question_id.strip('_')  # Remove leading/trailing underscores
                
                # Determine field type and attributes
                field_type, attributes = determine_field_type(text)
                
                question_data = {
                    "id": question_id,
                    "question": cleaned_text,
                    "type": field_type,
                    **attributes
                }
                
                sections[current_section]["questions"].append(question_data)
            
    return sections

def process_docx_files(docx_dir="../templates/docx_files", output_file="../data/training_data.json"):
    """Process all DOCX files and convert them to training data"""
    training_data = []
    
    for file in os.listdir(docx_dir):
        if file.endswith('.docx') and 'Questionnaire' not in file:
            doc_path = os.path.join(docx_dir, file)
            try:
                example = process_business_plan(doc_path)
                training_data.append(example)
                print(f"Processed: {file}")
            except Exception as e:
                print(f"Error processing {file}: {str(e)}")
    
    # Save training data
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(training_data, f, indent=2, ensure_ascii=False)
    
    print(f"\nProcessed {len(training_data)} business plans")
    return training_data

def convert_to_template():
    """Convert the E2 questionnaire to our template format"""
    doc_path = "../templates/docx_files/E2 Questionnaire.docx"
    questionnaire = extract_questionnaire_sections(doc_path)
    
    # Save the template
    with open("../templates/questionnaire_template.json", "w") as f:
        json.dump(questionnaire, f, indent=4)
    
    print("Generated questionnaire template from E2 Questionnaire.docx")

if __name__ == "__main__":
    # Process all business plan DOCX files
    process_docx_files()
    convert_to_template()